# 6
from docx import Document

document = Document()
document.save('yaheltest.docx')

# 6 Extra:
from docx import *
print (os.getcwd())
document = Document('yaheltest.docx')
document.add_heading("test")
document.add_paragraph("My first paragraph")
myparagraph = document.add_paragraph("Yahel test \n")
myparagraph.add_run ("my text")
document.save('yaheltest.docx')

Text = []
for para in document.paragraphs:
    Text.append(para.text)
    print((para.text))
#print(Text)
